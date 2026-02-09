from docx import Document
from docx.shared import Pt
import sys
import re

def reset_docx_middle(path):
    doc = Document(path)
    print(f"Cleaning TOC block in {path}")
    
    # 1. Encontrar la portada (sección 0)
    # El contenido real empieza después del primer salto de página manual o después del índice
    
    start_idx = -1
    placeholders = ["CONTENTS", "INDEX", "ÍNDICE", "TABLE OF CONTENTS"]
    
    for i, p in enumerate(doc.paragraphs):
        if any(ph == p.text.upper().strip() for ph in placeholders):
            start_idx = i
            break
    
    if start_idx == -1:
        print("No se encontró marcador de índice. Buscando el primer H1 real.")
        for i, p in enumerate(doc.paragraphs):
            if i > 5 and ('heading 1' in p.style.name.lower() or 'heading1' in p.style.name.lower()):
                start_idx = 1 # Empezamos a limpiar desde arriba
                break

    if start_idx != -1:
        # Borrar todo desde start_idx hasta el primer H1 real que NO sea el índice
        # o hasta encontrar una página que parezca contenido real (basado en Texto Largo o H1)
        paragraphs_to_remove = []
        for i in range(start_idx, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            p_text = p.text.strip()
            
            # Si encontramos 1. Strategic Proposal (el inicio real del doc)
            if "STRATEGIC PROPOSAL" in p_text.upper() and len(p_text) < 100:
                break
            
            # Si es el placeholder o parece parte del índice broken (líneas con números de página al final)
            paragraphs_to_remove.append(p)
            
        print(f"Removing {len(paragraphs_to_remove)} paragraphs...")
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)
            
        # Insertar el placeholder básico
        doc.paragraphs[0].insert_paragraph_before("TABLE OF CONTENTS")
        
    doc.save(path)
    print("Done clearing.")

if __name__ == "__main__":
    reset_docx_middle(sys.argv[1])
