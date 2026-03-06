from docx import Document
from docx.shared import Inches, Pt
import os

def restore_cover(docx_path, img_path):
    doc = Document(docx_path)
    print(f"Restoring cover to {docx_path}")
    
    # 1. Crear una nueva sección al principio
    # Word no permite "insertar sección al principio" fácilmente, 
    # pero podemos manipular el XML o reconstruir.
    # Reconstrucción mínima es mejor:
    
    new_doc = Document()
    
    # SECCIÓN 0: PORTADA
    section = new_doc.sections[0]
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    
    para = new_doc.add_paragraph()
    run = para.add_run()
    run.add_picture(img_path, width=section.page_width, height=section.page_height)
    
    # SECCIÓN 1: EL RESTO
    new_sec = new_doc.add_section()
    new_sec.top_margin = Inches(1)
    new_sec.bottom_margin = Inches(1)
    new_sec.left_margin = Inches(1)
    new_sec.right_margin = Inches(1)
    
    # Copiar párrafos (exceptuando el placeholder que ya pondrá generate_toc)
    # Buscamos el inicio del contenido real
    content_started = False
    for p in doc.paragraphs:
        p_text = p.text.upper().strip()
        if p_text in ["CONTENTS", "INDEX", "ÍNDICE", "TABLE OF CONTENTS"]:
             # Insertamos el placeholder para que generate_toc lo encuentre
             new_p = new_doc.add_paragraph("TABLE OF CONTENTS")
             new_p.style = p.style
             content_started = True
             continue
        
        if content_started:
            new_p = new_doc.add_paragraph(p.text, style=p.style)
            # Copiar runs para mantener negritas/itálicas básicas si es posible
            # (Aunque add_paragraph(text, style) es más seguro para estilos puros)
            new_p.alignment = p.alignment
            
    new_doc.save(docx_path)
    print("Cover restored and doc rebuilt.")

if __name__ == "__main__":
    restore_cover(
        r"c:\Users\Usuario\Workspace\01_Proyectos\Anclora-Agents-Skills\triumph-renta-mallorca-SEO\entregables\AI_Consultancy_Proposal_EN.docx",
        r"c:\Users\Usuario\Workspace\01_Proyectos\Anclora-Agents-Skills\triumph-renta-mallorca-SEO\pro_covers\MASTERPIECE_V7_PROPUESTA_EN.jpg"
    )
