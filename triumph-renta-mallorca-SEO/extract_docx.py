from docx import Document
import os

def extract_docx(file_path, output_path):
    doc = Document(file_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# CONTENIDO EXTRAÍDO DEL ANEXO CONTRACTUAL\n\n")
        
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n\n")
        
        for table in doc.tables:
            f.write("--- TABLA ---\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_text) + "\n")
            f.write("\n")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python extract_docx.py <input_docx> <output_md>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2]
    extract_docx(docx_path, output_path)
