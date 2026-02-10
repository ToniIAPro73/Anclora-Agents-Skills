import os
import re
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    """
    Adds centered page numbering to a paragraph.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = paragraph.add_run()
    
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = ""
    page_run._r.append(t1)

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    page_run._r.append(fldChar1)

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    page_run._r.append(instrText)

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_run._r.append(fldChar2)

def parse_markdown(content):
    """
    Extracts headers and structure from Markdown.
    """
    lines = content.split('\n')
    headers = []
    
    header_pattern = re.compile(r'^(#{1,3})\s+(.*)$')
    # Simple list bullet in bold
    list_pattern = re.compile(r'^(?:\s*)(?:[\*\-\+]|\d+\.)\s+\*\*(.+?)\*\*(?:\s*:?.*)$')

    for line in lines:
        h_match = header_pattern.match(line)
        if h_match:
            depth = len(h_match.group(1))
            text = h_match.group(2).strip()
            text = re.sub(r'^(?:\d+(?:\.\d+)*|[A-Z])[\.\s\)-]+', '', text).strip()
            headers.append({'level': depth, 'text': text, 'type': 'heading'})
            continue
            
        # We only treat bold list items as part of the structure if they look like headers
        # But for TOC H1-H3, we stick mostly to H tags unless asked otherwise.
        # However, the user previous request implied bold list items are part of the structure.
        # User requested H1 to H3 for TOC specifically.

    return headers

def find_cover(filename, covers_dir):
    """
    Matches a filename to a cover image.
    """
    name_upper = filename.upper()
    lang = "ES"
    if "_EN" in name_upper: lang = "EN"
    elif "_DE" in name_upper: lang = "DE"
    
    doc_type = "PROPUESTA"
    if "ANEXO" in name_upper: doc_type = "ANEXO"
    elif "PLAN" in name_upper: doc_type = "PLAN"
    elif "RESUMEN" in name_upper: doc_type = "RESUMEN"
    elif "ROI" in name_upper: doc_type = "ROI"
    
    cover_name = f"{doc_type}_{lang}.jpg"
    cover_path = os.path.join(covers_dir, cover_name)
    
    if os.path.exists(cover_path):
        return cover_path
    
    # Fallback to ES if not found
    fallback = os.path.join(covers_dir, f"{doc_type}_ES.jpg")
    return fallback if os.path.exists(fallback) else None

def generate_pro_docx(input_md, output_docx, cover_path):
    doc = Document()
    
    # 1. PAGE 1: COVER
    if cover_path:
        section = doc.sections[0]
        # Remove margins for cover
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.page_height = Inches(11.69) # A4
        section.page_width = Inches(8.27)
        
        doc.add_picture(cover_path, width=Inches(8.27), height=Inches(11.69))
        doc.add_section() # Section break for TOC
    
    # 2. PAGE 2: TOC (CONTENIDO)
    toc_section = doc.sections[-1]
    toc_section.left_margin = Inches(1.5)
    toc_section.right_margin = Inches(1.5)
    toc_section.top_margin = Inches(1.0)
    
    # Enable page numbering in footer for this section
    footer = toc_section.footer
    add_page_number(footer.paragraphs[0])

    # Title "CONTENIDO"
    for _ in range(3): doc.add_paragraph()
    t_p = doc.add_paragraph()
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t_p.add_run("CONTENIDO")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    for _ in range(2): doc.add_paragraph()

    # Read content to get headers
    with open(input_md, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    headers = parse_markdown(md_text)
    
    # Simple hierarchical numbering
    counters = [0] * 4
    for h in headers:
        level = h['level']
        for i in range(level + 1, 4): counters[i] = 0
        counters[level] += 1
        num = ".".join(str(counters[i]) for i in range(1, level + 1))
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3 * (level - 1))
        entry = p.add_run(f"{num} {h['text']}")
        entry.font.size = Pt(11)
        if level == 1: entry.bold = True

    doc.add_page_break()

    # 3. PAGE 3+: CONTENT
    # We parse MD and add paragraphs. This is simplified; 
    # for a full conversion we'd use a MD parser to docx.
    # For now, let's just put the text in to represent the document.
    
    # Re-reading to process full content
    lines = md_text.split('\n')
    for line in lines:
        if line.startswith('#'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            p = doc.add_heading(text, level=min(level, 4))
        elif line.strip():
            doc.add_paragraph(line.strip())

    doc.save(output_docx)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--covers", required=True)
    args = parser.parse_args()

    if not os.path.exists(args.output):
        os.makedirs(args.output)

    files = [f for f in os.listdir(args.input) if f.endswith('.md')]
    for f in files:
        md_path = os.path.join(args.input, f)
        docx_name = os.path.splitext(f)[0] + ".docx"
        output_path = os.path.join(args.output, docx_name)
        
        cover = find_cover(f, args.covers)
        print(f"Generating: {docx_name} (Cover: {os.path.basename(cover) if cover else 'None'})")
        generate_pro_docx(md_path, output_path, cover)

if __name__ == "__main__":
    main()
