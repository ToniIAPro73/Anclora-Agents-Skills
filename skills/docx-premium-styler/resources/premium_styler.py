import re
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_spacing(doc):
    """Añade separación entre el encabezado y el texto inicial."""
    if doc.paragraphs:
        # Dar espacio al primer párrafo
        first_para = doc.paragraphs[0]
        first_para.paragraph_format.space_before = Pt(48)

def set_para_layout(para):
    """Configura el layout del párrafo: evita cortes y justifica el texto."""
    # Evitar cortes entre páginas
    para.paragraph_format.keep_together = True
    
    style = para.style
    is_heading = style and hasattr(style, 'name') and style.name and style.name.startswith('Heading')
    
    # Justificar texto (excepto títulos)
    if not is_heading:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 'keep with next' para títulos
    if is_heading:
        para.paragraph_format.keep_with_next = True

def clean_markdown_artifacts(doc):
    """Elimina restos de markdown como '---'."""
    to_remove = []
    for para in doc.paragraphs:
        # También eliminar líneas que solo contienen guiones bajos o asteriscos de separación
        txt = para.text.strip()
        if txt == "---" or txt == "***" or re.match(r'^[-*_]{3,}$', txt):
            to_remove.append(para)
    for p in to_remove:
        p_element = p._element
        p_element.getparent().remove(p_element)

def set_table_borders(table):
    """Añade bordes manuales si el estilo de tabla falla."""
    tbl_pr = table._element.xpath('w:tblPr')[0]
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tbl_pr.append(borders)

def style_tables_professionally(doc):
    """Aplica un estilo premium a las tablas detected."""
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except:
            set_table_borders(table)
            
        # Formatear la primera fila
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'EEEEEE') # Gris muy claro
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True

def convert_markdown_tables(doc):
    """Busca bloques de texto que parecen tablas markdown y los convierte."""
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        para = paras[i]
        text = para.text.strip()
        if text.startswith('|') and text.endswith('|'):
            table_data = []
            start_idx = i
            while i < len(paras) and (paras[i].text.strip().startswith('|') or re.match(r'^\|\s*:?-+:?\s*\|', paras[i].text.strip())):
                line = paras[i].text.strip()
                if not re.match(r'^\|\s*:?-+:?\s*\|', line):
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    if cells:
                        table_data.append(cells)
                i += 1
            
            if len(table_data) > 1:
                target_para = paras[start_idx]
                new_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for r_idx, r_cells in enumerate(table_data):
                    for c_idx, c_text in enumerate(r_cells):
                        if c_idx < len(new_table.columns):
                            new_table.cell(r_idx, c_idx).text = c_text
                
                target_para._element.addprevious(new_table._element)
                for j in range(start_idx, i):
                    p_to_del = paras[start_idx]
                    p_element = p_to_del._element
                    p_element.getparent().remove(p_element)
                    paras.pop(start_idx)
                i = start_idx
                continue
        i += 1

def add_heading_numbering(doc):
    """Añade numeración automática inteligente."""
    counters = [0] * 10
    for para in doc.paragraphs:
        style = para.style
        if not (style and hasattr(style, 'name') and style.name and style.name.startswith('Heading')):
            continue
            
        try:
            level = int(style.name.replace('Heading ', ''))
        except:
            continue
            
        counters[level] += 1
        for j in range(level + 1, 10):
            counters[j] = 0
            
        # Omitir el título de la portada si tiene Heading (heurística: primer párrafo heading)
        # Pero aquí asumimos que queremos numerar todo lo que tenga estilo Heading
        
        num_parts = [str(counters[j]) for j in range(1, level + 1)]
        num_str = ".".join(num_parts) + ". "
        
        # Limpiar números previos si existen
        current_text = para.text.strip()
        para.text = re.sub(r'^\d+(\.\d+)*\.\s*', '', current_text)
            
        if para.runs:
            para.runs[0].text = num_str + para.runs[0].text
        else:
            para.add_run(num_str)

def ensure_toc(doc):
    """Añade o asegura el campo de Tabla de Contenidos."""
    idx_para = None
    for para in doc.paragraphs:
        txt = para.text.upper()
        if "ÍNDICE" in txt or "CONTENTS" in txt or "TABLA DE CONTENIDO" in txt:
            idx_para = para
            break
            
    if idx_para:
        run = idx_para.add_run()
        run.add_break()
        run.add_break()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)

    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)

def apply_premium_style(input_path, output_path):
    print(f"Applying premium style to {input_path}...")
    doc = Document(input_path)
    # Orden importa: primero tablas (pueden crear nuevos párrafos/títulos)
    convert_markdown_tables(doc)
    add_heading_numbering(doc)
    style_tables_professionally(doc)
    clean_markdown_artifacts(doc)
    ensure_toc(doc)
    add_header_spacing(doc)
    for para in doc.paragraphs:
        set_para_layout(para)
    doc.save(output_path)
    return True

if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit(1)
    apply_premium_style(sys.argv[1], sys.argv[2])
