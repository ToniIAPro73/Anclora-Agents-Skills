import os
import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_headers(content):
    """
    Parses markdown content and returns a list of headers and bold list items.
    Treats bold list items as sub-headers.
    """
    lines = content.split('\n')
    extracted = []
    
    # Regex for # to ###### headers
    header_pattern = re.compile(r'^(#{1,6})\s+(.*)$')
    # List pattern for bold titles: 1. **Title**, - **Title**, * **Title**
    # We look for lines that start with a list marker and have the entire text (or start) in bold
    list_title_pattern = re.compile(r'^(\s*)(?:[\*\-\+]|\d+\.)\s+\*\*(.+?)\*\*(?:\s*:?.*)$')

    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        h_match = header_pattern.match(line)
        if h_match:
            depth = len(h_match.group(1))
            text = h_match.group(2).strip()
            # ONLY strip numbering that looks like "1.", "1.1", "A.", etc. at the start
            # This regex avoids stripping words like "Resumen"
            text = re.sub(r'^(?:\d+(?:\.\d+)*|[A-Z])[\.\s\)-]+', '', text).strip()
            extracted.append({'depth': depth, 'text': text})
            continue
            
        l_match = list_title_pattern.match(line)
        if l_match:
            indent = len(l_match.group(1))
            text = l_match.group(2).strip()
            # If the list item has more text after the bold part, we might want to include it or just the bold.
            # User example: 1.1.1 Barrera de Acceso para Clientes en el Móvil
            # In the MD: 1. **Barrera de Acceso para Clientes en el Móvil**
            # If there's more after the closing **, we stick to the bold part as the title.
            depth = 7 + (indent // 2)
            extracted.append({'depth': depth, 'text': text})

    if not extracted:
        return []

    # Level Normalization:
    # If a document uses [1, 3, 7], transform it to [1, 2, 3]
    # This ensures that H1 is level 1, H3 is level 2 if there's no H2, etc.
    unique_depths = sorted(list(set(item['depth'] for item in extracted)))
    depth_to_level = {old_depth: new_level + 1 for new_level, old_depth in enumerate(unique_depths)}
    
    for item in extracted:
        item['level'] = depth_to_level[item['depth']]
    
    return extracted

def generate_hierarchical_numbers(headers):
    """
    Assigns correct hierarchical numbers (1, 1.1, 1.1.1) to headers.
    """
    counters = [0] * 10 # Support up to 10 levels
    last_level = 0
    
    for h in headers:
        level = h['level']
        
        # Reset counters for deeper levels if we move up (e.g., from 3 to 2)
        if level < last_level:
            for i in range(level + 1, 10):
                counters[i] = 0
        
        counters[level] += 1
        
        # Build number string: e.g., 1.2.1
        num_parts = [str(counters[i]) for i in range(1, level + 1)]
        h['number'] = ".".join(num_parts)
        last_level = level
        
    return headers

def create_toc_docx(headers, output_path):
    """
    Creates a Word document with the TOC formatted according to specifications.
    """
    doc = Document()
    
    # Section margins
    section = doc.sections[0]
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # 1. Centered "CONTENIDO" title
    # Separated 3 lines from the top
    for _ in range(3):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("CONTENIDO")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'

    # 2. Separated 2 lines from the index
    for _ in range(2):
        doc.add_paragraph()

    # 3. Add Index entries
    for h in headers:
        p = doc.add_paragraph()
        # Custom indentation based on level (0.3 inches per level)
        p.paragraph_format.left_indent = Inches(0.3 * (h['level'] - 1))
        p.paragraph_format.space_after = Pt(2)
        
        # We use a non-breaking space after the number
        run = p.add_run(f"{h['number']} {h['text']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        # Emphasis for level 1
        if h['level'] == 1:
            run.bold = True

    doc.save(output_path)

def process_file(input_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    filename = os.path.basename(input_path)
    # Change .md to .docx and add to pro_indices
    output_filename = os.path.splitext(filename)[0] + "_INDICE.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    print(f"Processing {input_path} -> {output_path}")
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    headers = parse_markdown_headers(content)
    headers = generate_hierarchical_numbers(headers)
    create_toc_docx(headers, output_path)

def main():
    parser = argparse.ArgumentParser(description="Generate Word TOC from Markdown")
    parser.add_argument("--input", required=True, help="Input Markdown file or directory")
    parser.add_argument("--output_dir", default="pro_indices", help="Output directory")
    
    args = parser.parse_args()
    
    if os.path.isdir(args.input):
        for root, dirs, files in os.walk(args.input):
            for file in files:
                if file.endswith(".md"):
                    process_file(os.path.join(root, file), args.output_dir)
    else:
        process_file(args.input, args.output_dir)

if __name__ == "__main__":
    main()
