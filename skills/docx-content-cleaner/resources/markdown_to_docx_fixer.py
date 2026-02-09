import re
import sys
from docx import Document
from docx.shared import Pt
from pathlib import Path

def apply_markdown_to_paragraph(para):
    """
    Analiza el texto de un párrafo, busca patrones de markdown y lo reconstruye
    con el formato nativo de Word.
    """
    full_text = para.text
    if not full_text:
        return False

    # Patrones básicos
    # Bold: **text**
    # Italic: *text*
    # Links: [text](url)
    # Headings artifacts: ### Text (solo si al inicio)
    
    has_markdown = bool(re.search(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\)|^#+\s+)', full_text))
    
    if not has_markdown:
        return False

    # Guardar el estilo original antes de limpiar
    original_style = para.style
    
    # Procesar encabezados al inicio
    heading_match = re.match(r'^(#+)\s+(.*)', full_text)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        full_text = content
        # Asignar estilo de encabezado
        style_name = f'Heading {min(level, 9)}'
        try:
            para.style = style_name
        except:
            para.style = 'Heading 1' # Fallback

    # Limpiar el párrafo para reconstruirlo
    # Nota: Esto borra cualquier formato previo que no sea el del estilo de párrafo.
    for run in para.runs:
        run.text = ""
    
    # regex compleja para encontrar todos los tokens de markdown
    # Grupos: 1: Bold, 2: Italic, 3: Link text, 4: Link URL
    pattern = re.compile(r'\*\*(.*?)\*\*|\*(.*?)\*|\[(.*?)\]\((.*?)\)')
    
    last_idx = 0
    for match in pattern.finditer(full_text):
        # Texto normal antes del match
        if match.start() > last_idx:
            para.add_run(full_text[last_idx:match.start()])
            
        if match.group(1): # Bold
            para.add_run(match.group(1)).bold = True
        elif match.group(2): # Italic
            para.add_run(match.group(2)).italic = True
        elif match.group(3): # Link
            # Por simplicidad en este script, ponemos el link en azul/subrayado
            # Word maneja Hyperlinks de forma especial, pero esto es una buena aproximación
            run = para.add_run(f"{match.group(3)} ({match.group(4)})")
            run.font.color.rgb = None # Usar color de estilo o manual si se desea
            run.underline = True
            
        last_idx = match.end()
        
    # Resto del texto
    if last_idx < len(full_text):
        para.add_run(full_text[last_idx:])
        
    return True

def clean_docx(input_path, output_path):
    doc = Document(input_path)
    modified = False
    
    for para in doc.paragraphs:
        if apply_markdown_to_paragraph(para):
            modified = True
            
    # También procesar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if apply_markdown_to_paragraph(para):
                        modified = True
    
    if modified:
        doc.save(output_path)
        return True
    return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python markdown_to_docx_fixer.py <input.docx> <output.docx>")
        sys.exit(1)
        
    inp = sys.argv[1]
    out = sys.argv[2]
    
    if clean_docx(inp, out):
        print(f"Documento limpiado y guardado en: {out}")
    else:
        print("No se detectó formato markdown que limpiar.")
