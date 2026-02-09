import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER

def add_bookmark(paragraph, name):
    """Inserta un bookmark al inicio del párrafo."""
    tag = paragraph._p
    # Bookmark start
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), name)
    bm_start.set(qn('w:name'), f'head_{name}')
    tag.insert(0, bm_start)
    # Bookmark end
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), name)
    tag.append(bm_end)

def insert_pageref_field(paragraph, bm_id):
    """Inserta un campo PAGEREF para un bookmark."""
    run = paragraph.add_run("\t") # Tabulador para activar el dot leader
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = f' PAGEREF head_{bm_id} \\h '
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    # Placeholder
    resText = OxmlElement('w:t')
    resText.text = '1'
    run._r.append(resText)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def is_potential_header(para):
    """Detecta el nivel lógico del párrafo para numeración y espaciado."""
    text = para.text.strip()
    if not text or len(text) > 130: return None
    
    # Filtro de seguridad para evitar procesar el propio índice
    clean_comp = re.sub(r'^\d+(\.\d+)*\.?\s*', '', text).upper()
    if any(f in clean_comp for f in ["ÍNDICE", "CONTENTS", "TABLA DE CONTENIDOS"]):
        return None
        
    style = para.style.name.lower()
    
    # 1. Por Estilo (H1, H2, H3 generados)
    if style == 'heading 1': return 1
    if style == 'heading 2': return 2
    if style == 'heading 3': return 3 # Cambiamos temporalmente para espaciado universal
    
    # 2. Heurística para listas MD (1. **...**) -> Nivel 3 para numeración
    if re.match(r'^\d+[\.\)]\s+', text) and any(run.bold for run in para.runs):
        return 3
            
    return None

def add_page_number_to_footer(paragraph):
    """Inserta el campo PAGE en el párrafo."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Inicia el campo
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    # Especifica que es el campo PAGE
    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "
    run._r.append(instrText)
    
    # Separador
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    # Placeholder para el visor (Word lo actualizará)
    t = OxmlElement('w:t')
    t.text = "1"
    run._r.append(t)
    
    # Fin del campo
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def get_style_safely(doc, style_name):
    """Busca un estilo por nombre o por su ID interno para evitar KeyErrors."""
    try: return doc.styles[style_name]
    except KeyError:
        # Buscar por ID (sin espacios) o por coincidencia de nombre
        search_name = style_name.replace(" ", "")
        for s in doc.styles:
            if s.name == style_name or s.name.replace(" ", "") == search_name:
                return s
    return None

def generate_toc_in_doc(input_path, output_path, renumber=True, lang='es'):
    print(f"Refinando Layout, Márgenes y Paginación (V16-Final): {input_path}")
    doc = Document(input_path)
    
    # --- 1. CONFIGURACIÓN DE MÁRGENES ---
    content_sec = doc.sections[0]
    for s in doc.sections:
        if s.left_margin > Pt(20):
            content_sec = s
            break
    usable_width = content_sec.page_width - content_sec.left_margin - content_sec.right_margin

    # --- 2. LOCALIZACIÓN Y LIMPIEZA DEL BLOQUE DE ÍNDICE ---
    # Lo hacemos ANTES de recolectar cabeceras para no indexar el índice anterior
    start_idx = -1
    placeholders = ["ÍNDICE DE CONTENIDOS", "TABLE OF CONTENTS", "INDEX", "ÍNDICE", "CONTENTS"]
    for i, para in enumerate(doc.paragraphs):
        p_text = para.text.upper().strip()
        if (any(ph == p_text for ph in placeholders) or "ÍNDICE" in p_text or "CONTENTS" in p_text) and len(p_text) < 40:
            start_idx = i
            break
    
    if start_idx == -1: 
        start_idx = 1
        insert_para = doc.paragraphs[1]
    else:
        # Limpieza SEGURA del bloque de índice previo
        paragraphs_to_remove = []
        for i in range(start_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            p_xml = p._element.xml
            
            # Si hay un salto de página, es el final de nuestro índice
            if '<w:br w:type="page"/>' in p_xml or '<w:lastRenderedPageBreak/>' in p_xml:
                paragraphs_to_remove.append(p)
                break
                
            # Si encontramos el inicio del contenido real (cualquier Heading)
            p_style = p.style.name.lower()
            if 'heading' in p_style or 'título' in p_style:
                break
                
            # Límite de seguridad para no entrar en bucle infinito o borrar todo el documento
            if i > start_idx + 30: 
                 break
                 
            paragraphs_to_remove.append(p)
            
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)
        insert_para = doc.paragraphs[start_idx]

    # --- 3. RECOLECCIÓN DE CABECERAS Y NUMERACIÓN ---
    def get_num_level(para):
        s = para.style.name.lower()
        if 'heading 1' in s or 'heading1' in s: return 1
        if 'heading 2' in s or 'heading2' in s: return 2
        if 'heading 3' in s or 'heading3' in s: return 3
        if re.match(r'^\d+[\.\)]\s+', para.text.strip()) and any(r.bold for r in para.runs): return 4
        return None

    header_data = []
    counters = [0] * 10
    
    for para in doc.paragraphs:
        # Evitar detectar el placeholder del índice como una cabecera para el propio índice
        if para == insert_para: continue
        
        lvl = get_num_level(para)
        if lvl and lvl <= 3:
            header_data.append((para, lvl))
            if renumber:
                counters[lvl] += 1
                for j in range(lvl + 1, 10): counters[j] = 0
                num_str = ".".join([str(counters[k]) for k in range(1, lvl + 1) if counters[k] > 0]) + ". "
                for r in para.runs:
                    r.text = re.sub(r'^\d+(\.\d+)*\.?\s*', '', r.text)
                if para.runs: para.runs[0].text = num_str + para.runs[0].text
                else: para.text = num_str

    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        if 'heading' in para.style.name.lower():
            if i + 1 < len(doc.paragraphs):
                nxt = doc.paragraphs[i+1]
                if nxt.text.strip() != "" and 'heading' not in nxt.style.name.lower():
                    nxt.insert_paragraph_before("")

    # --- 4. CONSTRUCCIÓN DEL NUEVO ÍNDICE ---
    insert_para.text = ""
    toc_title = "ÍNDICE" if lang == 'es' else "CONTENTS"
    
    h1_style = get_style_safely(doc, 'Heading 1')
    title = insert_para.insert_paragraph_before(toc_title)
    if h1_style: title.style = h1_style
    
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    
    body_style = get_style_safely(doc, 'Normal') or get_style_safely(doc, 'Body Text')
    
    bm_id = 7000
    for para, lvl in header_data:
        add_bookmark(para, str(bm_id))
        entry = insert_para.insert_paragraph_before("")
        if body_style: entry.style = body_style
        
        entry.paragraph_format.left_indent = Pt(12 * (lvl - 1))
        entry.paragraph_format.space_after = Pt(2)
        tab_stops = entry.paragraph_format.tab_stops
        tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        
        clean_text = para.text.strip().split('\n')[0]
        run = entry.add_run(clean_text)
        run.font.size = Pt(10)
        
        if lvl == 1: run.bold = True
        else: run.font.color.rgb = RGBColor(60, 60, 60)
            
        insert_pageref_field(entry, str(bm_id))
        bm_id += 1

    # Asegurar salto de página tras el índice si no existe
    p_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p == insert_para: p_idx = idx; break
        
    if p_idx != -1 and p_idx + 1 < len(doc.paragraphs):
        next_p = doc.paragraphs[p_idx+1]
        if '<w:br w:type="page"/>' not in next_p._element.xml:
            insert_para.insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)
    
    insert_para._element.getparent().remove(insert_para._element)

    # --- 5. NUMERACIÓN DE PÁGINAS ---
    # La Sección 0 es la portada y NO debe tener pie de página
    # La Sección 1 es el Índice, donde empieza la numeración visualmente (aunque sea pág 1)
    
    for i, section in enumerate(doc.sections):
        # Asegurar que la portada no tiene pie
        if i == 0:
            section.footer.is_linked_to_previous = False
            for p in section.footer.paragraphs:
                p.text = ""
            continue
            
        footer = section.footer
        
        # Desvincular pie de página de la sección 1 (Índice) respecto a la 0 (Portada)
        if i == 1:
            footer.is_linked_to_previous = False
            
        section.footer_distance = Pt(36)
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = "" 
        
        # Centrar numeración
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_to_footer(p)

    # --- 5. FINALIZAR ---
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        settings.append(OxmlElement('w:updateFields'))
    
    doc.save(output_path)
    print(f"Éxito: Documento guardado en {output_path} (renumber={renumber}, lang={lang})")

if __name__ == "__main__":
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        
        # Valores por defecto basados en el nombre
        filename = input_file.upper()
        lang = 'en' if ('_EN' in filename or 'PROPOSAL' in filename) else 'es'
        renum = False if ('_EN' in filename or 'PROPOSAL' in filename) else True
        
        # Overrides por argumentos
        for arg in sys.argv[3:]:
            if arg.startswith('renumber='):
                renum = arg.split('=')[1].lower() == 'true'
            elif arg.startswith('lang='):
                lang = arg.split('=')[1]
                
        print(f"Generando TOC: {input_file} -> {output_file} (renumber={renum}, lang={lang})")
        generate_toc_in_doc(input_file, output_file, renumber=renum, lang=lang)
    else:
        print("Uso: python generate_toc.py <input.docx> <output.docx>")
