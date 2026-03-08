import os
import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# Configuración de Rutas
COVERS_DIR = Path(r"c:\Users\Usuario\Workspace\01_Proyectos\Anclora-Agents-Skills\triumph-renta-mallorca-SEO\pro_covers")

def get_cover_path(md_filename):
    """Selecciona la portada correcta basada en el nombre del archivo markdown."""
    name = md_filename.lower()
    lang = "ES"
    if "_en" in name: lang = "EN"
    elif "_de" in name: lang = "DE"
    
    doc_type = "RESUMEN" 
    if "plan" in name or "ejecucion" in name or "acelerado" in name: doc_type = "PLAN"
    elif "propuesta" in name or "proposal" in name: doc_type = "PROPUESTA"
    elif "anexo" in name: doc_type = "ANEXO"
    elif "roi" in name: doc_type = "ROI"
    
    pattern = f"MASTERPIECE_V7_{doc_type}_{lang}"
    for f in COVERS_DIR.glob("*.jpg"):
        if pattern.upper() in f.name.upper():
            return f
    return COVERS_DIR / f"MASTERPIECE_V7_{doc_type}_ES.jpg"

def add_full_page_cover(doc, img_path):
    """Añade la portada como imagen a tamaño completo."""
    section = doc.sections[0]
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(str(img_path), width=section.page_width, height=section.page_height)
    
    # El section break ya actúa como salto de página
    new_section = doc.add_section()
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)

def ensure_toc(doc):
    """Añade un marcador para el Índice que el skill analyzer reemplazará."""
    p = doc.add_paragraph("ÍNDICE DE CONTENIDOS")
    p.style = 'Heading 2'

def set_para_layout(para, is_heading=False):
    """Configuración premium: Justificado y evitar cortes."""
    para.paragraph_format.keep_together = True
    if is_heading:
        para.paragraph_format.keep_with_next = True
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def parse_md(md_path, doc):
    """Convierte el md a la estructura del doc con numbering y tablas."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    counters = [0] * 10
    in_table = False
    table_data = []
    
    # Espaciado inicial tras cabecera (heurística: primera línea tras TOC)
    doc.add_paragraph().paragraph_format.space_before = Pt(36)

    for line in lines:
        line = line.strip()
        if not line or line in ["---", "***"]: continue
            
        # Detección de títulos
        h_match = re.match(r'^(#+)\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).replace('**', '')
            
            # Limpiar numeración previa en el texto (e.g. "1. Título" -> "Título")
            text = re.sub(r'^\d+(\.\d+)*\.?\s*', '', text)
            
            # Gestionar contadores de niveles
            counters[level] += 1
            for j in range(level + 1, 10): counters[j] = 0
            
            # Construir num_str evitando ceros si saltamos niveles
            active_parts = []
            for j in range(1, level + 1):
                if counters[j] > 0:
                    active_parts.append(str(counters[j]))
                else:
                    # Si hay un nivel intermedio en 0, lo ponemos como 1 para evitar 1.0.1
                    counters[j] = 1
                    active_parts.append("1")
                    
            num_str = ".".join(active_parts) + ". "
            
            p = doc.add_paragraph(num_str + text)
            try:
                p.style = f'Heading {min(level, 9)}'
            except:
                pass
            set_para_layout(p, True)
            continue
            
        # Tablas
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            if re.match(r'^\|\s*:?-+:?\s*\|', line): continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells: table_data.append(cells)
            continue
        else:
            if in_table and table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, r_cells in enumerate(table_data):
                    for c_idx, c_text in enumerate(r_cells):
                        if c_idx < len(table.columns):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = c_text
                            if r_idx == 0:
                                for tp in cell.paragraphs:
                                    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in tp.runs: run.bold = True
                in_table = False
                table_data = []
        
        # Listas
        li_match = re.match(r'^[\*\-\+]\s+(.*)', line)
        if li_match:
            p = doc.add_paragraph(li_match.group(1), style='List Bullet')
            set_para_layout(p)
            continue
            
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        set_para_layout(p)

def process_all_md(source_dir, output_dir):
    s_path = Path(source_dir)
    o_path = Path(output_dir)
    o_path.mkdir(exist_ok=True)
    
    for f in s_path.glob("*.md"):
        out_name = f.stem + ".docx"
        target = o_path / out_name
        print(f"Generating: {f.name} -> {out_name}")
        doc = Document()
        
        cover = get_cover_path(f.name)
        if cover and cover.exists(): add_full_page_cover(doc, cover)
        
        ensure_toc(doc)
        parse_md(f, doc)
        doc.save(target)

if __name__ == "__main__":
    src = r"c:\Users\Usuario\Workspace\01_Proyectos\Anclora-Agents-Skills\triumph-renta-mallorca-SEO\Doc_md"
    out = r"c:\Users\Usuario\Workspace\01_Proyectos\Anclora-Agents-Skills\triumph-renta-mallorca-SEO\entregables"
    process_all_md(src, out)
